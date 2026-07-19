from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(r"C:\hub-git\output\김기원_백엔드개발자_이력서_최종.docx")
PHOTO = Path(r"C:\Users\kshkj\OneDrive\바탕 화면\증명사진_김기원.png")

LATIN_FONT = "Calibri"
KOREAN_FONT = "Malgun Gothic"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1D2733"
MUTED = "5F6B7A"
LIGHT = "F4F6F9"
LIGHT_BLUE = "E8EEF5"
BORDER = "D7DEE7"
WHITE = "FFFFFF"
GREEN = "146C43"

# compact_reference_guide preset with named resume overrides:
# A4 page, Korean font fallback, compact 10.2 pt body, and 0.70 in side margins.
PAGE_WIDTH_DXA = 11906
MARGIN_DXA = 1008
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - (MARGIN_DXA * 2)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def set_table_geometry(table, widths_dxa, indent_dxa=0):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")

    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            none = {"val": "nil"}
            set_cell_border(cell, top=none, start=none, bottom=none, end=none, insideH=none, insideV=none)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name=LATIN_FONT):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size=10.2, color=INK, bold=False):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=False, size=9.2):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), LATIN_FONT)
    rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rPr.append(rFonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abs_id = max(existing_abs, default=0) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abs_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), "bullet")
    lvl.append(numFmt)
    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), "•")
    lvl.append(lvlText)
    lvlJc = OxmlElement("w:lvlJc")
    lvlJc.set(qn("w:val"), "left")
    lvl.append(lvlJc)
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "420")
    tabs.append(tab)
    pPr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "420")
    ind.set(qn("w:hanging"), "230")
    pPr.append(ind)
    lvl.append(pPr)
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    rPr.append(color)
    lvl.append(rPr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abs_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)


def add_bullet(container, text, num_id, size=9.8, after=3.0, color=INK):
    p = container.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.21
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, size=8.5, color=MUTED)


def add_section_title(doc, text, subtitle=None):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    if subtitle:
        r = p.add_run(f"  {subtitle}")
        set_run_font(r, size=9.0, color=MUTED, bold=False)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B8CBE0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_subheading(doc, title, meta=None):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, size=11.9, color=NAVY, bold=True)
    if meta:
        m = p.add_run(f"  |  {meta}")
        set_run_font(m, size=9.2, color=MUTED, bold=False)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "single")
    start.set(qn("w:sz"), "12")
    start.set(qn("w:space"), "6")
    start.set(qn("w:color"), BLUE)
    pBdr.append(start)
    pPr.append(pBdr)
    return p


def add_label_value(doc, label, value, after=2.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.16
    a = p.add_run(label)
    set_run_font(a, size=9.7, color=DARK_BLUE, bold=True)
    b = p.add_run(value)
    set_run_font(b, size=9.7, color=INK)
    return p


def shade_paragraph(paragraph, fill=LIGHT, border_color=BORDER):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "single")
    start.set(qn("w:sz"), "16")
    start.set(qn("w:space"), "5")
    start.set(qn("w:color"), border_color)
    pBdr.append(start)
    pPr.append(pBdr)


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.63)
    section.bottom_margin = Inches(0.63)
    section.left_margin = Inches(0.70)
    section.right_margin = Inches(0.70)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4.2)
    normal.paragraph_format.line_spacing = 1.22

    h1 = styles["Heading 1"]
    h1.font.name = LATIN_FONT
    h1._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = LATIN_FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    h2.font.size = Pt(11.9)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(NAVY)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = LATIN_FONT
    h3._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    h3.font.size = Pt(10.3)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(5)
    h3.paragraph_format.space_after = Pt(2)

    bullet_num_id = create_bullet_numbering(doc)

    # Quiet running header/footer; first page title block carries the identity.
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "BACKEND DEVELOPER RESUME"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    set_paragraph_font(hp, size=7.8, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.35))
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    left = fp.add_run("김기원 | Backend Developer")
    set_run_font(left, size=8.3, color=MUTED)
    fp.add_run("\t")
    page_label = fp.add_run("Page ")
    set_run_font(page_label, size=8.3, color=MUTED)
    add_page_field(fp)

    # PAGE 1 — profile
    profile = doc.add_table(rows=1, cols=2)
    set_table_geometry(profile, [7550, CONTENT_WIDTH_DXA - 7550], indent_dxa=0)
    remove_table_borders(profile)
    left_cell, photo_cell = profile.rows[0].cells
    set_cell_margins(left_cell, top=0, start=0, bottom=40, end=160)
    set_cell_margins(photo_cell, top=0, start=100, bottom=40, end=0)
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    photo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = left_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("김기원")
    set_run_font(r, size=25, color=NAVY, bold=True)
    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("BACKEND DEVELOPER")
    set_run_font(r, size=10.5, color=BLUE, bold=True)
    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("010-8390-5599   |   kjjk3250@gmail.com   |   1996.02.17")
    set_run_font(r, size=9.4, color=MUTED)
    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_hyperlink(p, "GitHub  github.com/kimlony/hub", "https://github.com/kimlony/hub", size=9.1)
    r = p.add_run("   |   ")
    set_run_font(r, size=9.1, color=MUTED)
    add_hyperlink(p, "Notion  Easy Hub", "https://app.notion.com/p/BizBee-HUB-36cbaf605146817d8835c1b94dfd2b81#7ad7b21b3ec44d399faac1c0996d7d2", size=9.1)

    photo_p = photo_cell.paragraphs[0]
    photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    photo_p.paragraph_format.space_after = Pt(0)
    photo_shape = photo_p.add_run().add_picture(str(PHOTO), width=Inches(1.28), height=Inches(1.28))
    photo_shape._inline.docPr.set("descr", "김기원 증명사진")
    photo_shape._inline.docPr.set("title", "프로필 사진")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.22
    r = p.add_run("ERP·외부 시스템 연동과 주문수집 자동화에 강점을 가진 백엔드 개발자입니다. ")
    set_run_font(r, size=11.2, color=NAVY, bold=True)
    r = p.add_run("동기 API 기반 수집 과정에서 겪은 실패 추적과 데이터 유실 문제를 Kafka, Outbox Pattern, Retry/DLQ, Recovery 구조로 확장해 해결하고, 대량 데이터 처리 성능과 운영 안정성을 함께 개선해 왔습니다.")
    set_run_font(r, size=10.0, color=INK)
    shade_paragraph(p, fill=LIGHT, border_color=BLUE)

    add_section_title(doc, "핵심 역량")
    for item in [
        "Java·Spring Boot 기반 ERP 물류/영업 기능과 REST API 연동 개발",
        "외부 쇼핑몰·물류 주문수집 및 Node.js/Puppeteer 기반 웹 자동화",
        "Kafka·Outbox Pattern 기반 비동기 Job 처리와 Retry/DLQ/Recovery 설계",
        "Oracle SQL 튜닝·페이징 조회·CSV 변환을 통한 대량 데이터 처리 개선",
        "requestId, Lease·Fencing Token, DB Lock을 활용한 실패 추적·중복 방지·안전한 재처리",
        "Testcontainers·GitHub Actions·Docker·AWS EC2 기반 테스트 및 배포 자동화",
    ]:
        add_bullet(doc, item, bullet_num_id, size=9.7, after=2.2)

    add_section_title(doc, "기술 스택")
    skills = doc.add_table(rows=4, cols=2)
    set_table_geometry(skills, [1700, CONTENT_WIDTH_DXA - 1700], indent_dxa=120)
    labels = ["Backend", "Data / Messaging", "Automation / UI", "DevOps / Test"]
    values = [
        "Java · Spring Boot · REST API · MyBatis",
        "Oracle · PostgreSQL · Apache Kafka · Redis",
        "Node.js · Puppeteer · JavaScript · TypeScript · AngularJS · React",
        "Docker · Testcontainers · GitHub Actions · Git · SVN · AWS EC2",
    ]
    for i, row in enumerate(skills.rows):
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_shading(row.cells[1], WHITE)
        edge = {"val": "single", "sz": "4", "color": BORDER}
        for c in row.cells:
            set_cell_border(c, top=edge, start=edge, bottom=edge, end=edge)
            set_cell_margins(c, top=75, bottom=75, start=115, end=115)
        row.cells[0].text = labels[i]
        row.cells[1].text = values[i]
        set_paragraph_font(row.cells[0].paragraphs[0], size=8.8, color=DARK_BLUE, bold=True)
        set_paragraph_font(row.cells[1].paragraphs[0], size=9.1, color=INK)

    add_section_title(doc, "경력 및 학력 요약", "총 경력 4년 9개월")
    summary = doc.add_table(rows=1, cols=3)
    set_table_geometry(summary, [3300, 2590, CONTENT_WIDTH_DXA - 5890], indent_dxa=120)
    data = [
        ("주식회사 비즈비", "2023.07 - 재직 중", "백엔드 개발자 / 사원"),
        ("진우프린트 주식회사", "2020.10 - 2022.05", "제작관리 / 주임"),
        ("광운대학교", "2015.03 - 2020.02", "경영학과 / 졸업"),
    ]
    for row_data in data:
        cells = summary.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value
    header_cells = summary.rows[0].cells
    for i, title in enumerate(("회사 / 학교", "기간", "직무 / 전공")):
        header_cells[i].text = title
        set_cell_shading(header_cells[i], NAVY)
        set_paragraph_font(header_cells[i].paragraphs[0], size=8.8, color=WHITE, bold=True)
    set_repeat_table_header(summary.rows[0])
    for ri, row in enumerate(summary.rows):
        for ci, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=78, bottom=78, start=110, end=110)
            edge = {"val": "single", "sz": "4", "color": BORDER}
            set_cell_border(cell, top=edge, start=edge, bottom=edge, end=edge)
            if ri > 0:
                set_paragraph_font(cell.paragraphs[0], size=8.9, color=INK, bold=(ci == 0))

    # PAGE 2 — professional experience
    doc.add_page_break()
    add_section_title(doc, "경력 상세")
    add_subheading(doc, "주식회사 비즈비", "2023.07 - 재직 중 | 정규직 | 백엔드 개발자·사원")
    add_label_value(doc, "담당 업무  ", "ERP 및 외부 시스템 연동, 쇼핑몰 주문수집, 물류·영업 기능 개발, 대량 조회 및 엑셀 다운로드 개선")

    highlight = doc.add_paragraph()
    highlight.paragraph_format.left_indent = Inches(0.12)
    highlight.paragraph_format.right_indent = Inches(0.08)
    highlight.paragraph_format.space_before = Pt(3)
    highlight.paragraph_format.space_after = Pt(7)
    a = highlight.add_run("핵심 성과  ")
    set_run_font(a, size=9.8, color=DARK_BLUE, bold=True)
    b = highlight.add_run("100만 건+ 조회 안정화  ·  50만 건+ 다운로드 개선  ·  API 미제공 거래처 수집 자동화")
    set_run_font(b, size=9.8, color=INK, bold=True)
    shade_paragraph(highlight, fill=LIGHT_BLUE, border_color=BLUE)

    add_subheading(doc, "1. 외부 시스템 연동 및 주문수집 기능 개발", "2023.07 - 현재")
    for item in [
        "외부 쇼핑몰·물류 시스템의 주문·입고·출고 데이터를 REST API로 수신하고, 데이터 정합성 검증 후 ERP DB에 저장",
        "API를 제공하지 않는 거래처 몰의 로그인·메뉴 이동·엑셀 다운로드를 Node.js/Puppeteer로 자동화하고 ERP 형식으로 변환",
        "사방넷 설치 프로그램이 필요한 주문수집·송장 송신을 위해 Windows 기반 AWS EC2에 Node.js Worker 환경 구성 및 배포",
    ]:
        add_bullet(doc, item, bullet_num_id, size=9.7, after=3.0)
    add_label_value(doc, "사용 기술  ", "Java, Oracle, MyBatis, AngularJS, REST API, Node.js, Puppeteer, AWS EC2", after=4)

    add_subheading(doc, "2. 대용량 조회 Timeout 개선 및 현황 프로그램 개발", "2023.07 - 2026.06")
    for item in [
        "ERP 물류·영업 모듈의 대량 데이터 조회 Timeout을 분석하고 Oracle SQL 튜닝 및 페이징 조회 구조로 개선",
        "검색 조건과 사용자별 환경설정을 적용해 100만 건 이상 데이터를 보유한 환경에서도 안정적으로 조회할 수 있도록 개선",
        "주문 데이터 50만 건 이상을 그리드에서 엑셀로 내려받을 때 발생하던 처리 지연을 CSV 변환 방식으로 개선",
    ]:
        add_bullet(doc, item, bullet_num_id, size=9.7, after=3.0)
    add_label_value(doc, "사용 기술  ", "Java, Oracle, SQL, MyBatis, AngularJS", after=4)

    add_subheading(doc, "3. 프랜차이즈 물류·영업 ERP 기능 개발", "2023.07 - 2026.06")
    for item in [
        "사용자 주문 등록 시 재고 관련 데이터를 조회해 주문 가능 여부를 검증하는 비즈니스 로직 구현",
        "전표 처리 단계의 무조건 반올림으로 발생하던 금액 불일치를 해결하기 위해 회사별 소수점 처리 정책 기능 개발에 참여",
        "주문·출고·입고 데이터를 분석해 재고 및 물류 현황 화면을 개발하고, 백엔드 API부터 AngularJS 화면까지 전체 흐름 구현",
    ]:
        add_bullet(doc, item, bullet_num_id, size=9.7, after=3.0)

    add_subheading(doc, "진우프린트 주식회사", "2020.10 - 2022.05 | 정규직 | 제작관리·주임")
    for item in [
        "패키지 생산 공정, 물품 입출고, 다음 공정 진행 여부, 고객 대응 및 납기 일정 관리",
        "수작업에 의존하던 반복 확인 업무를 경험하며 시스템화의 필요성을 체감했고, 이를 계기로 개발자로 전향",
    ]:
        add_bullet(doc, item, bullet_num_id, size=9.7, after=3.0)

    # PAGE 3 — project
    doc.add_page_break()
    add_section_title(doc, "주요 프로젝트")
    add_subheading(doc, "Easy Hub — Kafka 기반 주문수집 자동화 플랫폼", "개인 프로젝트 | 진행 중")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("REST API 기반 동기 주문수집에서 경험한 실패 추적, 재처리, 중복 처리, 완료 여부 확인의 어려움을 해결하기 위해 만든 프로젝트입니다. ")
    set_run_font(r, size=10.4, color=INK)
    r = p.add_run("주문수집 요청을 Job으로 관리하고 Kafka Worker 기반 비동기 파이프라인으로 확장했습니다.")
    set_run_font(r, size=10.4, color=NAVY, bold=True)

    arch = doc.add_paragraph()
    arch.paragraph_format.left_indent = Inches(0.16)
    arch.paragraph_format.right_indent = Inches(0.10)
    arch.paragraph_format.space_before = Pt(3)
    arch.paragraph_format.space_after = Pt(5)
    r = arch.add_run("Hub API  →  hub_job / hub_job_outbox  →  Kafka  →  Worker  →  ERP·외부 시스템")
    set_run_font(r, size=9.7, color=DARK_BLUE, bold=True)
    shade_paragraph(arch, fill=LIGHT_BLUE, border_color=BLUE)

    add_subheading(doc, "설계 및 구현")
    for item in [
        "Spring Boot 기반 Hub API와 Node.js Worker 구조 설계, 주문수집 요청을 hub_job으로 관리하고 상태를 QUEUED·PROCESSING·SUCCESS·FAILED로 추적",
        "DB 저장과 Kafka 발행 사이의 메시지 유실 가능성을 줄이기 위해 Outbox Pattern 적용: hub_job과 hub_job_outbox를 같은 트랜잭션으로 저장하고 Publisher가 Kafka로 발행",
        "Kafka 기반 비동기 Job 처리와 DB Lock·Kafka Key를 활용한 동시성 제어로 동일 계정의 중복 수집 방지",
        "Worker 장애 후 중단 상태를 복구하는 Recovery Scanner와 반복 실패 작업을 격리하는 Retry/DLQ 구조 구현",
        "requestId 기준으로 Kafka lag, Worker 처리량, Job 상태, DLQ 발생 여부를 확인할 수 있는 모니터링 화면 구현",
        "로그인 계정의 회사 식별자를 기준으로 주문·Job·로그·ERP 전송 데이터의 접근 범위를 분리",
        "Testcontainers 기반 PostgreSQL/Kafka 통합 테스트와 GitHub Actions 기반 CI/CD 파이프라인 구성",
    ]:
        add_bullet(doc, item, bullet_num_id, size=9.6, after=2.6)

    add_subheading(doc, "ERP 연동 및 사용자 흐름")
    for item in [
        "주문수집, 정규화, ERP 변환, 수집 전송, 외부 API 조회, 엑셀 다운로드까지 하나의 흐름으로 추적",
        "자동 반영·수동 전송·외부 API 조회·엑셀 다운로드 이력을 동일 기준으로 확인할 수 있도록 상태 및 결과 기록 구조 설계",
        "회사 업무에서 경험한 동기 주문수집 구조를 실패 추적·재처리·중복 방지·운영 모니터링까지 포함한 비동기 구조로 확장",
    ]:
        add_bullet(doc, item, bullet_num_id, size=9.6, after=2.6)

    add_subheading(doc, "성과 및 검증")
    metrics = doc.add_table(rows=2, cols=3)
    set_table_geometry(metrics, [3297, 3297, CONTENT_WIDTH_DXA - 6594], indent_dxa=120)
    metric_values = [("100,000건", "주문수집 E2E 부하 테스트"), ("3.1배", "4 worker 처리량 개선"), ("1,214초 → 388초", "동일 조건 처리시간 단축")]
    for i, (big, label) in enumerate(metric_values):
        cell = metrics.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        cell.text = big
        set_paragraph_font(cell.paragraphs[0], size=14, color=WHITE, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = metrics.rows[1].cells[i]
        set_cell_shading(cell, LIGHT)
        cell.text = label
        set_paragraph_font(cell.paragraphs[0], size=8.8, color=DARK_BLUE, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in metrics.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=100, bottom=100, start=100, end=100)
            edge = {"val": "single", "sz": "4", "color": BORDER}
            set_cell_border(cell, top=edge, start=edge, bottom=edge, end=edge)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("테스트 확인  ")
    set_run_font(r, size=9.2, color=DARK_BLUE, bold=True)
    r = p.add_run("Kafka lag 0, DLQ 실패 메시지 없음, 주문수집·정규화·ERP 변환·전송·외부 API 조회·엑셀 다운로드 이력의 일관된 추적")
    set_run_font(r, size=9.2, color=INK)

    add_label_value(doc, "기술 스택  ", "Java, Spring Boot, PostgreSQL, Kafka, Node.js, TypeScript, Docker, Testcontainers, GitHub Actions, React", after=3)
    links = doc.add_paragraph()
    links.paragraph_format.left_indent = Inches(0.12)
    links.paragraph_format.right_indent = Inches(0.08)
    links.paragraph_format.space_before = Pt(3)
    links.paragraph_format.space_after = Pt(0)
    links.paragraph_format.line_spacing = 1.08
    r = links.add_run("서비스  ")
    set_run_font(r, size=8.6, color=DARK_BLUE, bold=True)
    add_hyperlink(links, "http://hub.rony.kr/", "http://hub.rony.kr/", size=8.6)
    r = links.add_run("    |    GitHub  ")
    set_run_font(r, size=8.6, color=DARK_BLUE, bold=True)
    add_hyperlink(links, "https://github.com/kimlony/hub", "https://github.com/kimlony/hub", size=8.6)
    r = links.add_run()
    r.add_break()
    set_run_font(r, size=8.6, color=INK)
    r = links.add_run("Notion  ")
    set_run_font(r, size=8.6, color=DARK_BLUE, bold=True)
    add_hyperlink(
        links,
        "https://app.notion.com/p/BizBee-HUB-36cbaf605146817d8835c1b94dfd2b81",
        "https://app.notion.com/p/BizBee-HUB-36cbaf605146817d8835c1b94dfd2b81#7ad7b21b3ec44d399faac1c0996d7d2",
        size=8.6,
    )
    shade_paragraph(links, fill=LIGHT, border_color=BLUE)

    # PAGE 4 — self introduction. Use page-break-before to avoid a blank page
    # when the project page ends exactly at the printable boundary.
    self_title = add_section_title(doc, "자기소개")
    self_title.paragraph_format.page_break_before = True

    add_subheading(doc, "개발자가 되고자 했던 이유")
    paragraphs = [
        "이전에는 패키지 제작관리 업무를 하며 생산 공정, 납기, 고객 대응을 담당했습니다. 당시 바코드 확인, 물품 입출고 확인, 다음 공정 진행 여부 확인처럼 반복적으로 확인해야 하는 업무가 많았고, 대부분 사람의 경험과 수작업에 의존하고 있었습니다.",
        "반복 업무를 시스템으로 관리하면 누락을 줄이고 더 안정적으로 운영할 수 있겠다는 생각이 개발자로 전향한 계기가 되었습니다. 단순히 화면에 기능을 추가하는 데 그치지 않고, 실제 업무 흐름을 이해해 사람이 반복하던 일을 시스템으로 바꾸는 개발을 지향합니다.",
    ]
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.30
        set_paragraph_font(p, size=10.3, color=INK)

    add_subheading(doc, "관심 있는 개발 분야")
    paragraphs = [
        "외부 시스템 연동, 데이터 수집 자동화, 대량 데이터 처리, 운영 안정성을 높이는 백엔드 구조에 관심이 있습니다. ERP와 외부 쇼핑몰·물류 시스템을 연동하며 외부 API가 항상 정상 응답을 주지 않고, 지연·실패·중복 요청·Timeout이 일상적으로 발생한다는 점을 경험했습니다.",
        "이후 Easy Hub에서 주문수집 요청을 Job으로 관리하고 Kafka Worker 기반 비동기 처리, Outbox Pattern, Retry/DLQ, Recovery, DB Lock을 적용했습니다. 앞으로도 사용자에게는 단순해 보이지만 내부적으로는 실패를 추적하고 안전하게 복구할 수 있는 시스템을 설계하는 개발자로 성장하고 싶습니다.",
    ]
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.30
        set_paragraph_font(p, size=10.3, color=INK)

    add_subheading(doc, "기억에 남는 개발 프로젝트")
    paragraphs = [
        "가장 기억에 남는 프로젝트는 개인 프로젝트 Easy Hub입니다. 회사에서 REST API 기반 동기 주문수집 기능을 개발하며 외부 API 실패, 데이터 유실 가능성, 로그 신뢰성, 처리 완료 여부 추적의 어려움을 경험했고, 이를 직접 해결해 보기 위해 프로젝트를 시작했습니다.",
        "주문수집 요청을 Job으로 관리하고 Kafka 기반 비동기 파이프라인으로 확장했으며, Outbox Pattern으로 DB 저장과 메시지 발행의 일관성을 높였습니다. Worker 장애 복구, Retry/DLQ, 동시성 제어, 모니터링까지 구현하고 Mock Mall 기반 100,000건 E2E 부하 테스트에서 4 worker 처리량이 1 worker 대비 약 3.1배 개선되는 것을 확인했습니다.",
        "현재는 AWS EC2에 운영 배포를 완료하고 Nginx와 도메인을 연결해 실제 운영 주소에서 접속할 수 있도록 구성했습니다. 또한 수집 Job별 Attempt 이력을 생성하고 종료 처리하는 구조를 구현해, 진행 중인 수집 과정과 각 시도의 완료·실패 여부를 관측할 수 있도록 개선했습니다.",
        "이 프로젝트를 통해 안정적인 백엔드는 비동기 처리 구조뿐 아니라 운영 환경에서 상태를 관측하고 문제 원인을 추적할 수 있어야 한다는 점을 배웠습니다. 앞으로도 외부 시스템이 불안정한 상황에서 실패를 안전하게 복구하고, 사용자가 신뢰할 수 있는 결과를 제공하는 백엔드 개발자가 되겠습니다.",
    ]
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.30
        set_paragraph_font(p, size=10.3, color=INK)

    # Document properties and compatibility-friendly defaults.
    props = doc.core_properties
    props.title = "김기원 백엔드 개발자 이력서"
    props.subject = "Backend Developer Resume"
    props.author = "김기원"
    props.keywords = "Java, Spring Boot, Kafka, ERP, Backend"
    settings = doc.settings.element
    even_odd = settings.find(qn("w:evenAndOddHeaders"))
    if even_odd is not None:
        settings.remove(even_odd)
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
