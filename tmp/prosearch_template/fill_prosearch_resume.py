from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run
from docx.shared import Pt, RGBColor


SOURCE = Path(r"C:\hub-git\tmp\prosearch_template\프로써치_양식_작업사본.docx")
OUTPUT = Path(r"C:\hub-git\output\지원포지션명_김기원(1996년생)_프로써치_추천.docx")


def set_run_font(run, *, size=11, bold=False, color=None, underline=False):
    run.font.name = "맑은 고딕"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_text(paragraph, text, *, size=11, bold=False, color=None, align=None):
    if align is not None:
        paragraph.alignment = align

    # 글자 노드만 바꾸고, 같은 문단에 연결된 선/텍스트 상자/VML 도형은 보존한다.
    text_nodes = paragraph._p.xpath("./w:r/w:t | ./w:hyperlink/w:r/w:t")
    run = None
    if text_nodes:
        text_nodes[0].text = text
        for node in text_nodes[1:]:
            node.text = ""
        run = Run(text_nodes[0].getparent(), paragraph)
    elif text:
        run = paragraph.add_run(text)
    if run is not None:
        set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def copy_paragraph_properties(target, source):
    if target._p.pPr is not None:
        target._p.remove(target._p.pPr)
    if source._p.pPr is not None:
        target._p.insert(0, deepcopy(source._p.pPr))


def add_hyperlink(paragraph, text, url, *, size=9.5):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "맑은 고딕")
    fonts.set(qn("w:hAnsi"), "맑은 고딕")
    fonts.set(qn("w:eastAsia"), "맑은 고딕")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(int(size * 2)))
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(int(size * 2)))
    run_properties.extend([fonts, color, underline, size_element, size_cs])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(run_properties)
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document(SOURCE)
p = doc.paragraphs

# 원본 양식의 페이지 시작 위치를 명시적으로 고정한다.
p[35].paragraph_format.page_break_before = True
p[71].paragraph_format.page_break_before = True
p[107].paragraph_format.page_break_before = True

# 1페이지: 기본사항
set_text(p[3], "지원분야 : 백엔드 개발자", size=12, bold=True, color=(0, 0, 255))
set_text(p[6], "성    명 : 김기원")
set_text(p[7], "생년월일 : 1996.02.17")
set_text(p[8], "결혼여부 :")
set_text(p[9], "주    소 :")
set_text(p[10], "현 직 장 : 주식회사 비즈비")
set_text(p[11], "직    급 : 사원")
set_text(p[12], "총  급여 :")
set_text(p[13], "희망급여 :")
set_text(p[14], "이직가능시기 :")
set_text(p[15], "연락처 / 이메일 : 010-8390-5599 / kjjk3250@gmail.com", size=10.5)

set_text(p[18], "2015.03 ~ 2020.02     광운대학교 경영학과 졸업")
set_text(p[19], "")
set_text(p[20], "")
set_text(p[23], "")

set_text(p[26], "▣ Java·Spring Boot 기반 ERP 물류·영업 기능 및 REST API 연동 개발", size=10.5)
set_text(p[27], "▣ Kafka·Outbox Pattern 기반 비동기 Job 처리와 Retry/DLQ/Recovery 설계", size=10.5)
set_text(p[28], "▣ Oracle SQL 튜닝·페이징·CSV 변환을 통한 대량 데이터 처리 개선", size=10.5)
set_text(p[29], "▣ AWS EC2·Docker·Testcontainers·GitHub Actions 기반 배포 및 테스트", size=10.5)

set_text(p[31], "경력사항 (총경력 4년 9개월)", size=12, bold=True, color=(0, 0, 255))
set_text(p[32], "2023.07 ~ 현재       주식회사 비즈비  백엔드 개발자·사원 재직 중", size=10.5)
set_text(p[33], "2020.10 ~ 2022.05     진우프린트 주식회사  제작관리·주임", size=10.5)

# 2페이지: 기술/경력기술서 첫 번째 회사
set_text(p[36], "외국어 :")
set_text(p[37], "컴퓨터 : Java, Spring Boot, Oracle, PostgreSQL, Kafka, Node.js, Docker, AWS EC2", size=10)
set_text(p[40], "")
set_text(p[43], "")
set_text(p[46], "")

set_text(
    p[50],
    "2023.07 ~ 현재     주식회사 비즈비  백엔드 개발자·사원 (3년 1개월)",
    size=11.5,
    bold=True,
    color=(0, 0, 255),
)
set_text(p[52], "☞ 사업개요 : 기업용 ERP 솔루션 개발·운영", size=10.5)
set_text(p[53], "☞ 업    종 : 정보통신업 / 소프트웨어 개발", size=10.5)
set_text(p[54], "☞ 매 출 액 :", size=10.5)
set_text(p[55], "☞ 직 원 수 :", size=10.5)
set_text(p[56], "☞ 주력사업 : ERP 및 외부 시스템 연동, 물류·영업 기능 개발", size=10.5)
set_text(p[59], "  ERP 및 외부 시스템 연동 개발", size=10.5, bold=True)
set_text(p[60], "- REST API 기반 주문·입고·출고 데이터 수집 및 ERP DB 저장", size=10.5)
set_text(p[61], "- Node.js/Puppeteer 기반 주문수집 자동화와 AWS EC2 배포", size=10.5)
set_text(p[64], "☞ 100만 건 이상 조회 환경의 Timeout을 SQL 튜닝·페이징으로 개선", size=10.5)
set_text(p[65], "☞ 50만 건 이상 다운로드 지연을 CSV 변환 방식으로 개선", size=10.5)
set_text(p[67], "[이직희망사유]", size=10.5)

# 3페이지: 두 번째 회사
set_text(
    p[71],
    "2020.10 ~ 2022.05     진우프린트 주식회사  제작관리·주임 (1년 8개월)",
    size=11.5,
    bold=True,
    color=(0, 0, 255),
)
set_text(p[73], "☞ 사업개요 : 패키지 제작·생산", size=10.5)
set_text(p[74], "☞ 업    종 : 인쇄·제작", size=10.5)
set_text(p[75], "☞ 매 출 액 :", size=10.5)
set_text(p[76], "☞ 직 원 수 :", size=10.5)
set_text(p[77], "☞ 주력사업 : 패키지 생산 및 물품 입출고 관리", size=10.5)
set_text(p[80], "☞ 제작관리·공정 및 납기 관리", size=10.5, bold=True)
set_text(p[81], "- 생산 공정, 물품 입출고, 다음 공정 진행 여부 관리", size=10.5)
set_text(p[82], "- 고객 대응과 납기 일정 조율", size=10.5)
set_text(p[85], "☞ 반복 확인 업무의 누락을 줄이기 위한 시스템화 필요성 발견", size=10.5)
set_text(p[86], "☞ 현장 업무 흐름과 사용자 관점의 요구사항 이해", size=10.5)
set_text(p[88], "[이직사유] 개발자로 직무 전환", size=10.5)

# 3페이지 하단의 기존 여백 문단을 활용한 개인 프로젝트
copy_paragraph_properties(p[90], p[71])
p[90].paragraph_format.page_break_before = False
set_text(
    p[90],
    "개인 프로젝트  Easy Hub — Kafka 기반 주문수집 자동화 플랫폼",
    size=11.5,
    bold=True,
    color=(0, 0, 255),
)
copy_paragraph_properties(p[91], p[72])
set_text(p[91], "▶ 프로젝트 개요", size=10.5, bold=True)
set_text(p[92], "☞ 동기 REST API 수집의 실패 추적·재처리·중복 처리 문제를 해결한 프로젝트", size=9.5)
set_text(p[93], "☞ 주문수집 요청을 Job으로 관리하고 Kafka Worker 비동기 파이프라인으로 확장", size=9.5)
set_text(p[94], "")
copy_paragraph_properties(p[95], p[79])
set_text(p[95], "▶ 주요 구현", size=10.5, bold=True)
set_text(p[96], "- Outbox Pattern으로 DB 저장과 Kafka 메시지 발행의 일관성 확보", size=9.5)
set_text(p[97], "- Retry/DLQ·Recovery Scanner·DB Lock으로 실패 복구와 중복 수집 방지", size=9.5)
set_text(p[98], "- AWS EC2 운영 배포, Nginx·도메인 연결, Job Attempt 이력·종료 처리", size=9.5)
set_text(p[99], "- Testcontainers 통합 테스트와 GitHub Actions CI/CD 구성", size=9.5)
copy_paragraph_properties(p[100], p[84])
set_text(p[100], "▶ 성과 및 운영", size=10.5, bold=True)
set_text(p[101], "☞ 100,000건 E2E 테스트: 4 worker 처리량이 1 worker 대비 약 3.1배 개선", size=9.5)
set_text(p[102], "☞ Kafka lag 0·DLQ 실패 없음 확인, 진행 중 수집과 Attempt 완료·실패 관측", size=9.5)
set_text(p[103], "기술: Java, Spring Boot, PostgreSQL, Kafka, Node.js, TypeScript, Docker, React", size=9)

set_text(p[104], "")
prefix = p[104].add_run("서비스 ")
set_run_font(prefix, size=9.5, bold=True)
add_hyperlink(p[104], "http://hub.rony.kr/", "http://hub.rony.kr/", size=9.5)
middle = p[104].add_run("   GitHub ")
set_run_font(middle, size=9.5, bold=True)
add_hyperlink(p[104], "https://github.com/kimlony/hub", "https://github.com/kimlony/hub", size=9.5)

set_text(p[105], "")
prefix = p[105].add_run("Notion ")
set_run_font(prefix, size=9, bold=True)
notion = "https://app.notion.com/p/BizBee-HUB-36cbaf605146817d8835c1b94dfd2b81"
add_hyperlink(p[105], notion, notion, size=9)
set_text(p[106], "")

# 4페이지: 자기소개서
set_text(p[110], "반복 확인에 의존하던 제작관리 업무를 시스템화하고 싶어 개발자로 전향했습니다.", size=10.5)
set_text(p[111], "ERP와 외부 시스템 연동 경험을 바탕으로 업무의 누락과 비효율을 줄이는 백엔드를 만들고자 합니다.", size=10.5)
set_text(p[112], "")
set_text(p[113], "")
set_text(p[114], "")

set_text(p[116], "외부 API의 지연·실패·중복을 기능 문제가 아닌 운영 문제로 보고 원인을 끝까지 추적합니다.", size=10.5)
set_text(p[117], "Kafka·Outbox·Retry/DLQ·Recovery를 적용해 실패를 확인하고 안전하게 복구할 수 있게 설계했습니다.", size=10.5)
set_text(p[118], "")
set_text(p[119], "")
set_text(p[120], "")
set_text(p[121], "")

set_text(p[123], "새 기술의 도입보다 장애를 예측하고 실패 원인을 추적하며 복구 가능한 구조를 만드는 일을 우선합니다.", size=10.5)
set_text(p[124], "사용자에게 단순해 보여도 내부적으로 데이터 일관성과 관측 가능성을 갖춘 시스템을 지향합니다.", size=10.5)
set_text(p[125], "")
set_text(p[126], "")
set_text(p[127], "")
set_text(p[128], "")

set_text(p[130], "Easy Hub를 EC2에 운영 배포하고 Nginx·도메인 연결, Job Attempt 이력과 종료 처리를 구현했습니다.", size=10.5)
set_text(p[131], "외부 시스템이 불안정해도 실패를 안전하게 복구하고 신뢰할 수 있는 결과를 제공하는 개발자가 되겠습니다.", size=10.5)
set_text(p[132], "")
set_text(p[133], "")
set_text(p[134], "")
set_text(p[135], "")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
